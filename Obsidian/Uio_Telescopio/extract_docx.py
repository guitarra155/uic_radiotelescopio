import sys
from docx import Document

def extract_content(docx_path):
    doc = Document(docx_path)
    content = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text.strip())
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content.append(" | ".join(row_text))
                
    return '\n\n'.join(content)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <path_to_docx>")
        sys.exit(1)
    
    path = sys.argv[1]
    try:
        text = extract_content(path)
        print(text)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
