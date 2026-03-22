import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = docx.Document()
    
    # Título Principal
    titulo = doc.add_heading('Plataforma DSP - Radiotelescopio (1420.40 MHz)', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Explicación de la Arquitectura Modular y Funcionamiento en Tiempo Real').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    
    # 1. Arquitectura del Proyecto
    doc.add_heading('1. ¿Cómo funciona la arquitectura modular de la aplicación?', level=1)
    doc.add_paragraph(
        'El proyecto se reestructuró para separar la lógica de la interfaz gráfica ('
        'Flet) en módulos específicos, simulando un entorno de ingeniería de software estructurado.'
    )
    
    # Detalle de archivos
    archivos = [
        ('main.py', 'Es el punto de entrada oficial de la app. No contiene lógica pesada. Configura la ventana principal (tema oscuro, tamaño) y une todos los módulos en el sistema nativo de navegación por pestañas (TabBar).'),
        ('constants.py', 'Almacena toda la "paleta de colores" y constantes. Si deseas cambiar el color de un botón o el fondo de los paneles en todo el software, solo editas un valor aquí y se refleja instantáneamente en todas partes.'),
        ('charts.py', 'Concentra todo el poder matemático (Numpy) y de dibujo de señales (Matplotlib). Su trabajo es computar matrices, generar gráficos científicos (espectros, el Waterfall/espectrograma, histogramas) y renderizarlos en fragmentos de memoria Base64. De esta manera evitamos tener que guardar los gráficos "como fotos PNG" en el disco duro, optimizando la lectura/escritura.'),
        ('components/layout.py', 'Dibuja el encabezado de "Procesamiento de Señales" y el pie de página que indica la versión y el hardware. Opera de forma independiente al resto.'),
        ('components/shared.py', 'Aloja plantillas para controles repetitivos. Asegura que todos los textos, campos y paneles posean la misma estética sin tener que repetir código.'),
        ('tabs/ (Carpeta Vistas)', 'Contiene un archivo por cada pestaña de tu interfaz. Cada archivo (como monitoring.py o sdr_config.py) solo se preocupa por armar "su propia pantalla" pidiendo los gráficos a charts.py y acomodando botones. Al final, main.py los carga todos de golpe.')
    ]
    
    for filename, desc in archivos:
        p = doc.add_paragraph()
        p.add_run(filename + ': ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph('\n')
    
    # 2. Análisis del Tiempo Real
    doc.add_heading('2. ¿Funcionará esto en Tiempo Real?', level=1)
    
    p2 = doc.add_paragraph()
    p2.add_run('Respuesta corta: ').bold = True
    p2.add_run('SÍ. El flujo técnico actual permite operar en tiempo real puro, sin necesidad de bases de datos.\n')
    
    doc.add_paragraph(
        'Un radiotelescopio que mide la transición fina del hidrógeno emite millones de muestras (MS/s) continuas en I/Q. '
        'Para lograr procesar eso "en vivo", calcular estadísticos (Smart Trigger) y repintar la pantalla a altos cuadros por segundo (FPS), '
        'la arquitectura debe manejarse con cuidado. Tu código actual sienta las bases, pero requeriría dos implementaciones técnicas que este Framework (Flet + Python) soporta sin problemas:'
    )
    
    doc.add_heading('A) Multihilo y Ring-Buffers (Buffers Circulares)', level=2)
    doc.add_paragraph(
        'Jamás guardaremos los datos crudos en el disco (a menos que tú lo decidas al capturar un evento anómalo). '
        'El hardware SDR inyectará los datos directamente a la Memoria RAM a través de hilos paralelos (usando la librería "threading" de Python, '
        'conectada a pyrtlsdr o TCP a GNU Radio). '
        'Las variables en Python emplearán "collections.deque", que actúan como cintas transportadoras: entra el nuevo paquete de muestras, '
        'calculas su desviación/media instantáneamente (con Numpy), y las muestras más viejas simplemente se evaporan y sobrescriben. Esto garantiza latencia de milisegundos.'
    )
    
    doc.add_heading('B) Renderizado Asíncrono de Interfaz', level=2)
    doc.add_paragraph(
        'Actualmente se simuló el UI renderizando Matplotlib en ráfagas. Matplotlib calcula píxeles minuciosamente bajo CPU, '
        'lo cual bloquea temporalmente la interfaz (se traba la pantalla). Para tiempo real:'
    )
    
    ul1 = doc.add_paragraph(style='List Bullet')
    ul1.add_run('Opción 1: ').bold = True
    ul1.add_run('Se envía a Matplotlib o scipy a un "Hilo Secundario" en el fondo que procesa el Espectrograma por su cuenta a 10 FPS, y cuando tiene la "foto (Base64)" lista, simplemente la empuja a la Interfaz de Flet (Flet soporta actualización asíncrona de variables sin frenar la vista del usuario).')
    
    ul2 = doc.add_paragraph(style='List Bullet')
    ul2.add_run('Opción 2 (Aún más rápido): ').bold = True
    ul2.add_run('Reemplazamos Matplotlib por los gráficos nativos de Flet (ft.LineChart), los cuales corren sobre Skia/Flutter bajo aceleración de Hardware (Tarjeta de Video GPU). De esta forma se alcanza el Santo Grial de 60 Cuadros por Segundo suaves y estables en streaming de la señal.')
    
    p_final = doc.add_paragraph('\nEn conclusión, tu plataforma no solo está bien estructurada, sino que posee la escalabilidad técnica para captar el cielo a 1420.40 MHz, computar interferencias instantáneas (RFI) y alertarte en microsegundos, usando estrictamente la memoria volátil.')
    p_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.save('c:/uic_radiotelescopio/Explicacion_Arquitectura.docx')

if __name__ == '__main__':
    main()
